#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Apr 19 17:44:28 2024

@author: nishkarshgupta
"""

import pandas as pd
from docx import Document
from langchain_community.llms import Ollama
from collections import Counter

# Initialize the Llama2 model on Ollama
llm = Ollama(base_url='http://localhost:11434', model="llama2")

# Load your CSV data
reddit_data = pd.read_csv('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/CombinedMatt.csv')

# Combine ThreadTitle and MsgBody if both exist, otherwise use what's available
reddit_data['text'] = reddit_data.apply(
    lambda row: row['ThreadTitle'] + " " + row['MsgBody'] if pd.notnull(row['MsgBody']) else row['ThreadTitle'], axis=1)

# Create a new Word document for results
results_doc = Document()

# Function to classify text and add results to the Word document
def classify_text(text):
    response = llm.invoke(f"Classify the following text into categories: television, movies, books. Text: {text}")
    if response:
        results_doc.add_paragraph(f"Text: {text}\nClassification: {response.strip()}\n")
        return response.strip()
    else:
        error_msg = "Error: Unable to classify text"
        results_doc.add_paragraph(f"Text: {text}\nClassification: {error_msg}\n")
        return error_msg

# # Classify texts and add to document for the entire DataFrame
# for text in reddit_data['text']:
#     classify_text(text)

# Classify texts and add to document for the first 100 entries of the DataFrame using slicing
for text in reddit_data['text'][0:100]:
    classify_text(text)


# Save the document with classification results
results_doc.save('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/ClassificationResults_300.docx')

# Load the document to process counts
doc = Document('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/ClassificationResults_300.docx')

# Initialize counters for each category
television_shows = Counter()
movies = Counter()
books = Counter()

# Define functions to add to the counters
def add_to_counter(counter, item):
    if item:
        counter[item] += 1

# Process the document to fill counters
for para in doc.paragraphs:
    parts = para.text.split('*')
    for part in parts:
        if 'Television:' in part:
            show_name = part.split('"')[1] if '"' in part else None
            add_to_counter(television_shows, show_name)
        elif 'Movies:' in part:
            movie_name = part.split('"')[1] if '"' in part else None
            add_to_counter(movies, movie_name)
        elif 'Books:' in part:
            book_name = part.split('"')[1] if '"' in part else None
            add_to_counter(books, book_name)

# Create another document for counts
counts_doc = Document()

# Function to write results to document
def write_results_to_doc(counter, title, document):
    document.add_heading(title, level=1)
    for item, count in counter.items():
        document.add_paragraph(f'{item} - {count}')

# Write simplified results to the document
write_results_to_doc(television_shows, "Television Shows", counts_doc)
write_results_to_doc(movies, "Movies", counts_doc)
write_results_to_doc(books, "Books", counts_doc)

# Save the document with counts
counts_doc.save('/Users/nishkarshgupta/Desktop/ASU Coursework/Q4/Big Data/Files/CountsResults_300.docx')
